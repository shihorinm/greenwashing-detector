"""
Wordレポート生成モジュール
"""
from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

def generate_word_report(results: Dict[str, Any]) -> bytes:
    """
    診断結果からWordレポートを生成
    
    Args:
        results: 診断結果
        
    Returns:
        Wordデータ（バイト列）
    """
    doc = Document()
    
    # 日本語フォント設定
    set_japanese_font(doc)
    
    # タイトル
    title = doc.add_heading('ClimateWash診断レポート', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(46, 125, 50)
    
    # 診断日時
    doc.add_paragraph(f"診断日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
    
    # コンテンツタイプと診断対象
    content_type = results.get('content_type', '不明')
    doc.add_paragraph(f"コンテンツタイプ: {content_type}")
    
    # 診断対象コンテンツを追加
    content_sample = results.get('content_sample', '')
    if content_sample:
        doc.add_paragraph(f"診断対象: {content_sample[:500]}")
    
    doc.add_page_break()
    
    # サマリーセクション
    doc.add_heading('診断サマリー', level=2)
    
    # サマリーテーブル
    risk_level = results.get('overall_risk', '不明')
    score = results.get('score', 0)
    risk_info = results.get('risk_info', {})
    
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    # ヘッダー行
    summary_table.rows[0].cells[0].text = '項目'
    summary_table.rows[0].cells[1].text = '内容'
    
    # データ行
    summary_table.rows[1].cells[0].text = '総合評価'
    summary_table.rows[1].cells[1].text = risk_level
    
    summary_table.rows[2].cells[0].text = 'スコア'
    summary_table.rows[2].cells[1].text = f"{score}/100"
    
    summary_table.rows[3].cells[0].text = '適用指令'
    summary_table.rows[3].cells[1].text = results.get('directives', '不明')
    
    summary_table.rows[4].cells[0].text = '診断バージョン'
    summary_table.rows[4].cells[1].text = results.get('version', '不明')
    
    summary_table.rows[5].cells[0].text = '違反項目数'
    summary_table.rows[5].cells[1].text = f"{len(results.get('violations', []))}件"
    
    # 評価説明
    doc.add_paragraph()
    description = risk_info.get('description', '')
    p = doc.add_paragraph()
    p.add_run('評価: ').bold = True
    p.add_run(description)
    
    # まとめ
    doc.add_paragraph()
    doc.add_heading('総括', level=2)
    summary_text = results.get('summary', '')
    doc.add_paragraph(summary_text)
    
    doc.add_page_break()
    
    # 違反項目セクション
    violations = results.get('violations', [])
    if violations:
        doc.add_heading(f'検出された問題点 ({len(violations)}件)', level=2)
        
        for i, violation in enumerate(violations, 1):
            # 違反項目のタイトル
            v_title = f"{i}. {violation.get('category_name', '不明な項目')} (項目 {violation.get('category', '')})"
            p = doc.add_paragraph()
            p.add_run(v_title).bold = True
            
            # 詳細テーブル
            v_table = doc.add_table(rows=4, cols=2)
            v_table.style = 'Light List Accent 1'
            
            v_table.rows[0].cells[0].text = 'リスクレベル'
            v_table.rows[0].cells[1].text = violation.get('risk_level', 'Unknown')
            
            v_table.rows[1].cells[0].text = '減点'
            v_table.rows[1].cells[1].text = f"{violation.get('points_deducted', 0)}点"
            
            v_table.rows[2].cells[0].text = '問題内容'
            v_table.rows[2].cells[1].text = violation.get('description', '')
            
            v_table.rows[3].cells[0].text = '該当表現'
            v_table.rows[3].cells[1].text = violation.get('evidence', '')
            
            doc.add_paragraph()
    else:
        doc.add_heading('検出された問題点', level=2)
        doc.add_paragraph('問題は検出されませんでした。')
    
    doc.add_page_break()
    
    # 是正提案セクション
    recommendations = results.get('recommendations', [])
    if recommendations:
        doc.add_heading(f'是正提案 ({len(recommendations)}件)', level=2)
        
        for i, rec in enumerate(recommendations, 1):
            # 提案のタイトル
            r_title = f"{i}. {rec.get('issue', '問題')}"
            p = doc.add_paragraph()
            p.add_run(r_title).bold = True
            
            # 現在の表現
            p = doc.add_paragraph()
            p.add_run('現在の表現: ').bold = True
            p.add_run(f"「{rec.get('current_expression', '')}」")
            
            # 推奨する表現
            p = doc.add_paragraph()
            p.add_run('推奨する表現: ').bold = True
            p.add_run(f"「{rec.get('recommended_expression', '')}」")
            
            # 理由
            p = doc.add_paragraph()
            p.add_run('理由: ').bold = True
            p.add_run(rec.get('explanation', ''))
            
            doc.add_paragraph()
    else:
        doc.add_heading('是正提案', level=2)
        doc.add_paragraph('是正の必要はありません。')
    
    # メモリ上に保存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def set_japanese_font(doc):
    """文書全体に日本語フォントを設定"""
    # デフォルトスタイルに日本語フォントを設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Noto Sans CJK JP'
    font.size = Pt(11)
    
    # 東アジア言語のフォント設定
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK JP')
